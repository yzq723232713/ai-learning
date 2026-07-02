"""
Day 15-17 综合练习：文档加载 → 切块 → Embedding → 检索

你要实现完整的 RAG 离线建库 + 在线检索流程。
每个函数的逻辑自己写，注释已经给了详细的提示。
"""
import numpy as np
from pathlib import Path
from typing import List, Optional
from datetime import datetime
from pydantic import BaseModel
import csv
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from sentence_transformers import SentenceTransformer


# ============================================================
# 第1步：定义数据结构（来自 Day 15）
# ============================================================
# 提示：
#   - LoadedDocument: 加载后的文档，包含 content(字符串), source(文件名),
#     file_type(文件类型), page_count(页数)
#   - Chunk: 切分后的文本块，包含 text(块内容), index(序号),
#     source(来源文件名，从 LoadedDocument 继承)

class LoadedDocument(BaseModel):
    """从文件加载后的文档"""
    content: str
    source: str
    file_type: str
    page_count: int = 1
    loaded_at: datetime = None

    def model_post_init(self, __content):
        if self.loaded_at is None:
            self.loaded_at = datetime.now()

class Chunk(BaseModel):
    """切分后的文本块"""
    text: str
    index: int
    source: str = "" 
    char_start: int
    char_end: int


# ============================================================
# 第2步：文档加载器（来自 Day 15）
# ============================================================

class DocumentLoader:
    """
    根据文件扩展名自动选择解析方式，统一返回 LoadedDocument。

    支持的格式：.pdf, .docx, .txt, .csv

    用法：
        loader = DocumentLoader()
        doc = loader.load("E:/learn/day15/test_docs/产品说明.txt")
        # doc.content  → 文档正文
        # doc.source   → "产品说明.txt"
        # doc.file_type → "txt"
    """

    @staticmethod
    def load(file_path: str) -> LoadedDocument:
        """
        根据扩展名路由到对应的加载方法。
        提示：
        - 用 Path(file_path).suffix.lower() 获取扩展名
        - 用 if/elif 判断扩展名，调用对应的 _load_xxx 方法
        - 不支持的格式抛出 ValueError
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return DocumentLoader._load_pdf(path)
        elif suffix == ".docx":
            return DocumentLoader._load_docx(path)
        elif suffix == ".txt":
            return DocumentLoader._load_txt(path)
        elif suffix == ".csv":
            return DocumentLoader._load_csv(path)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")

    @staticmethod
    def _load_pdf(path: Path) -> LoadedDocument:
        """
        用 PyMuPDF (fitz) 加载 PDF。
        提示：
        - doc = fitz.open(path) 打开文档
        - for page in doc: 遍历每一页
        - page.get_text() 获取页面的文本内容
        - 在 close 之前用 doc.page_count 获取页数
        - 所有页的文本用 "\n".join() 拼接
        """
        doc = fitz.open(path)
        full_text = []
        for page in doc:
            full_text.append(page.get_text())
        page_count = doc.page_count
        doc.close()
        content = "\n".join(full_text).strip()
        return LoadedDocument(
            content=content,
            source=path.name,
            file_type="pdf",
            page_count=page_count
        )


    @staticmethod
    def _load_docx(path: Path) -> LoadedDocument:
        """
        用 python-docx 加载 Word 文档。
        提示：
        - doc = DocxDocument(path) 打开文档
        - for p in doc.paragraphs: 遍历段落
        - p.text 获取段落文本
        - 跳过空段落（if p.text.strip()）
        """
        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n".join(paragraphs)
        return LoadedDocument(
            content=content,
            source=path.name,
            file_type="docx",
            page_count=len(doc.paragraphs) // 30 + 1
        )

    @staticmethod
    def _load_txt(path: Path) -> LoadedDocument:
        """
        加载纯文本文件。
        提示：
        - path.read_text(encoding="utf-8") 读取全部内容
        - strip() 去掉首尾空白
        """
        content = path.read_text(encoding="utf-8").strip()
        return LoadedDocument(
            content=content,
            source=path.name,
            file_type="txt"
        )

    @staticmethod
    def _load_csv(path: Path) -> LoadedDocument:
        """
        加载 CSV，转为 Markdown 表格格式，方便 LLM 理解。
        提示：
        - 用 csv.DictReader 读取，拿到 rows（列表，每个元素是字典）
        - headers = list(rows[0].keys()) 取列名
        - 第一行: " | ".join(headers)
        - 第二行: " | ".join(["---"] * len(headers))  分隔线
        - 后续行: 每行的值用 " | ".join() 拼接
        - 全部行用 "\n".join() 合并
        """
        with open(path, "r", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            rows = list(reader)

        if not rows:
            return LoadedDocument(content="", source=path.name, file_type="csv")
        headers = list(rows[0].keys())
        lines = [" | ".join(headers)]
        lines.append(" | ".join(["---"] * len(headers)))
        for row in rows:
            lines.append(" | ".join(str(row[h]) for h in headers))

        content = "\n".join(lines)
        return LoadedDocument(
            content=content,
            source=path.name,
            file_type="csv"
        )

# ============================================================
# 第3步：文本切分（来自 Day 16）
# ============================================================

class TextSplitter:
    """
    把长文档切成小块，为 Embedding 做准备。

    用法：
        splitter = TextSplitter(chunk_size=300)
        chunks = splitter.split(doc.content, source=doc.source)
        # 返回 List[Chunk]，每个 Chunk 带有 source 来源标记
    """
    @staticmethod
    def split(text: str, source: str = "", chunk_size: int = 300, separators: List[str] = None) -> List[Chunk]:
        """
        先用段落（\\n\\n）切，太长的段落再用句子（。）切，还太长就硬切。

        提示：
        - 调用 self._split_by_separators(text)
        - 检查是否有超长片段（len > self.chunk_size），超长的硬切
        - 硬切: for i in range(0, len(text), self.chunk_size)
        - 每个块创建 Chunk 对象，index 从 0 递增，source 传给 Chunk
        """
        if separators is None:
            separators = ["\n\n", "\n", "。", " "]

        chunks = TextSplitter._split_by_separators(text, separators, chunk_size)

        result = []
        idx = 0
        for chunk in chunks:
            if len(chunk) <= chunk_size:
                result.append(Chunk(text=chunk, index=idx, source=source, char_start=0, char_end=0))
                idx += 1
            else:
                for i in range(0, len(chunk), chunk_size):
                    result.append(Chunk(text=chunk[i: i + chunk_size], index=idx, source=source, char_start=0, char_end=0))
                    idx += 1
        
        return result
    
    @staticmethod
    def _split_by_separators(text: str, separators: List[str], max_size: int) -> List[str]:
        """
        按分隔符优先级切分：先 \\n\\n → \\n → 。→ 空格 → 放弃治疗。

        提示：
        - 定义一个 separators 列表: ["\\n\\n", "\\n", "。", " "]
        - 依次尝试每个分隔符
        - 如果分隔符不在文本中，跳过
        - 如果切出来的片段都 <= chunk_size，返回
        - 如果有超长片段，对那个片段用下一个分隔符递归切
        """
        if not separators:
            return [text]
        
        # 选取第一个切分符号
        sep = separators[0]
        
        # 判断text中是否存在这个分隔符，不存在则递归调用_split_by_separators，传入separators[1:]
        if sep not in text:
            return TextSplitter._split_by_separators(text, separators[1:], max_size)
        # 对文本切分，去空
        pieces = text.split(sep)
        pieces = [ p.strip() for p in pieces if p.strip()]

        # 检查是否所有片段都不超过限制，如果满足直接返回
        if all(len(p) <= max_size for p in pieces):
            return pieces
        # 存在超长chunk，递归调用_split_by_separators，传入separators[1:]
        result = []
        for p in pieces:
            if len(p) <= max_size:
                result.append(p)
            else:
                result.extend(TextSplitter._split_by_separators(p, separators[1:], max_size))

        return result

# ============================================================
# 第4步：Embedding + 向量检索（来自 Day 17）
# ============================================================

class VectorStore:
    """
    把 Chunk 转向量，支持相似度检索。

    用法：
        store = VectorStore(model_name="BAAI/bge-small-zh-v1.5")
        store.add_chunks(chunks)          # 批量向量化
        results = store.search("怎么安装?") # 检索 Top-5
    """

    def __init__(self, model_name: str = "BAAI/bge-small-zh-v1.5"):
        """
        加载 Embedding 模型。
        提示：
        - self.model = SentenceTransformer(model_name)
        - 首次运行会下载模型（已缓存则秒加载）
        """
        self.model = SentenceTransformer(model_name)
        self.embeddings = []
        self.chunks = []

    def add_chunks(self, chunks: List[Chunk]):
        """
        把所有块的文本转向量，存起来。
        提示：
        - texts = [c.text for c in chunks]  提取所有文本
        - self.embeddings = self.model.encode(texts)  批量向量化
        - self.chunks = chunks  保存块引用，检索时取回原文
        """
        texts = [c.text for c in chunks]
        self.embeddings.extend(self.model.encode(texts))
        self.chunks.extend(chunks)

    def search(self, query: str, top_k: int = 5) -> List[tuple]:
        """
        输入查询文本，返回最相关的 Top-K 个块。

        返回格式：[(相似度, Chunk对象), ...]，按相似度降序排列

        提示：
        - query_emb = self.model.encode(query)  查询转向量
        - 遍历所有块的向量，用余弦相似度计算分数
        - 余弦相似度: np.dot(v1,v2) / (norm(v1)*norm(v2))
        - 用 sorted() 按分数降序排列，取前 top_k 个
        """
        query_emb = self.model.encode(query)
        scores = []
        for i, (chunk, chunk_emb) in enumerate(zip(self.chunks, self.embeddings)):
            score = self.cosine_similarity(query_emb, chunk_emb)
            scores.append((score, i, chunk))

        scores.sort(reverse=True, key=lambda x:x[0])
        return scores[:top_k]

    @staticmethod
    def cosine_similarity(v1: np.ndarray, v2: np.ndarray) -> float:
        """计算两个向量的余弦相似度"""
        return np.dot(v1, v2) / (np.linalg.norm(v1) * np.linalg.norm(v2))


# ============================================================
# 第5步：组合成完整流水线
# ============================================================

class RAGPipeline:
    """
    一条命令完成：加载文档 → 切块 → 向量化。

    用法：
        pipeline = RAGPipeline()
        pipeline.ingest("E:/learn/day15/test_docs")  # 离线建库
        results = pipeline.query("怎么安装?")          # 在线检索
    """

    def __init__(self):
        self.store = VectorStore()

    def ingest(self, folder_path: str):
        """
        扫描文件夹中的所有文档，加载 → 切块 → 向量化。

        提示：
        - 用 Path(folder_path).iterdir() 遍历文件夹
        - 只处理 .pdf .docx .txt .csv 文件
        - 每个文件: loader.load() → splitter.split() → 收集所有块
        - 最后 store.add_chunks(all_chunks)
        - 打印每步的进度信息
        """
        allowed = {".pdf", ".docx", ".txt", ".csv"}
        for p in Path(folder_path).iterdir():
            if p.suffix.lower() not in allowed:
                continue
            ld = DocumentLoader.load(p)
            print(f"已加载文件：{p.name}")
            chunks = TextSplitter.split(ld.content, ld.source)
            print(f"已切块文件：{p.name}")
            self.store.add_chunks(chunks)
            print(f"已向量化文件：{p.name}")
            

    def query(self, question: str, top_k: int = 5):
        """
        检索与问题最相关的文档块。
        提示：直接调用 self.store.search(question, top_k)
        """
        return self.store.search(question,top_k)


# ============================================================
# 测试代码（你的实现让下面全部跑通）
# ============================================================
if __name__ == "__main__":
    import os
    os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"

    print("=" * 60)
    print("测试1: 文档加载")
    print("=" * 60)

    test_files = [
        "E:/learn/day15/test_docs/产品说明.txt",
        "E:/learn/day15/test_docs/会议纪要.pdf",
        "E:/learn/day15/test_docs/员工信息.csv",
        "E:/learn/day15/test_docs/需求文档.docx",
    ]

    for f in test_files:
        doc = DocumentLoader.load(f)
        preview = doc.content[:60].replace("\n", " ")
        print(f"  {doc.source:15s} | {doc.file_type:4s} | {len(doc.content):4d}字 | {preview}...")

    print(f"\n{'='*60}")
    print("测试2: 文本切分")
    print("=" * 60)

    doc = DocumentLoader.load("E:/learn/day15/test_docs/会议纪要.pdf")
    chunks = TextSplitter.split(doc.content, doc.source)

    for c in chunks:
        preview = c.text[:50].replace("\n", " ")
        print(f"  块{c.index}: {len(c.text):3d}字 | 来源:{doc.source} | {preview}...")

    print(f"\n{'='*60}")
    print("测试3: Embedding + 检索")
    print("=" * 60)

    store = VectorStore()
    store.add_chunks(chunks)

    queries = ["会议主题是什么？", "有哪些参会人？", "什么决议？"]
    for q in queries:
        results = store.search(q, top_k=3)
        print(f"\n  查询: 「{q}」")
        for rank, (score, _, chunk) in enumerate(results, 1):
            bar = "█" * int(score * 20)
            preview = chunk.text[:60].replace("\n", " ")
            print(f"    {rank}. [{score:.4f}] {preview}  {bar}   引用自：{chunk.source}")

    print(f"\n{'='*60}")
    print("测试4: 完整流水线")
    print("=" * 60)

    pipeline = RAGPipeline()
    pipeline.ingest("E:/learn/day15/test_docs")

    print("\n  检索测试:")
    for q in ["安装步骤", "保修期多久", "张三在哪个部门"]:
        results = pipeline.query(q, top_k=2)
        print(f"  「{q}」")
        for rank, (score, _, chunk) in enumerate(results, 1):
            preview = chunk.text[:80].replace("\n", " ")
            print(f"    {rank}. [{score:.4f}] {preview}  引用自：{chunk.source}")

    print("\n" + "=" * 60)
    print("全部测试通过!")
    print("=" * 60)
