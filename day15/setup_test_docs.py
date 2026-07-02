"""在 E:\learn\day15\test_docs 下生成三种格式的测试文档"""
from pathlib import Path

base = Path("E:/learn/day15/test_docs")
base.mkdir(exist_ok=True)

# 1. TXT
(base / "产品说明.txt").write_text(
    "产品名称: 智能温控器 X-100\n版本: 3.2\n安装步骤:\n"
    "1. 将设备固定在墙上\n2. 连接电源线\n3. 下载APP完成配对\n保修期: 2年\n",
    encoding="utf-8"
)

# 2. CSV
import csv
with open(base / "员工信息.csv", "w", newline="", encoding="utf-8-sig") as f:
    csv.writer(f).writerows([
        ["姓名", "部门", "入职日期", "薪资"],
        ["张三", "研发", "2022-01-15", "15000"],
        ["李四", "产品", "2021-06-20", "18000"],
        ["王五", "测试", "2023-03-10", "12000"],
    ])

# 3. PDF
from fpdf import FPDF
pdf = FPDF()
pdf.add_page()
pdf.add_font("SimSun", "", "C:/Windows/Fonts/simsun.ttc", uni=True)
pdf.set_font("SimSun", "", 14)
pdf.cell(0, 10, "会议纪要", align="C")
pdf.ln(15)
pdf.set_font("SimSun", "", 11)
pdf.multi_cell(0, 8,
    "会议主题: 2024年Q2产品规划\n"
    "时间: 2024年3月15日\n"
    "参会人: 张三(研发)、李四(产品)、王五(测试)\n\n"
    "议题一: 智能温控器迭代\n"
    "决议: 增加语音控制功能，预计5月底上线。\n\n"
    "议题二: APP界面优化\n"
    "决议: 首页增加快捷操作面板，6月初提测。"
)
pdf.output(str(base / "会议纪要.pdf"))
print(f"PDF 创建完成: {base / '会议纪要.pdf'}")

# 4. DOCX
from docx import Document
doc = Document()
doc.add_heading("需求规格说明书", level=1)
doc.add_paragraph("项目名称: 企业知识库系统", style="Normal")
doc.add_heading("功能需求", level=2)
doc.add_paragraph(
    "1. 支持 PDF、Word、文本文件上传\n"
    "2. 自动提取文本内容并建立索引\n"
    "3. 支持自然语言问答\n"
    "4. 答案附带来源引用"
)
doc.add_heading("非功能需求", level=2)
doc.add_paragraph(
    "响应时间 < 3秒\n"
    "支持并发用户数 >= 100\n"
    "数据安全: 文档加密存储"
)
doc.save(str(base / "需求文档.docx"))
print(f"DOCX 创建完成: {base / '需求文档.docx'}")

print("\n全部测试文档生成完毕:")
for f in sorted(base.iterdir()):
    size_kb = f.stat().st_size / 1024
    print(f"  {f.name:20s} {size_kb:8.1f} KB")
