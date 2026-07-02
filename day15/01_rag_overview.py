"""
Day 15：RAG 全景 + 文档加载器

RAG 全链路（必须能不看笔记画出来）：

  离线阶段（事前准备）:
    文档 → 加载 → 切块(Chunk) → Embedding → 存入向量库

  在线阶段（用户提问时）:
    问题 → Embedding → 向量库检索 → 拼接 Prompt → LLM 生成回答

今天做第1步：文档加载 — 把 PDF/Word/TXT/CSV 统一转为标准格式
"""
from pathlib import Path
from typing import List, Optional
from pydantic import BaseModel
from datetime import datetime
import csv
import fitz  # PyMuPDF
from docx import Document as DocxDocument


# ========================================
# Part 1: 理解 RAG 全链路
# ========================================

print("=== Part 1: RAG 全链路 ===\n")

print("""
                    RAG 系统全链路

    ┌─────────────────────────────────────────┐
    │            离线阶段（做一次）              │
    │                                          │
    │  PDF ─┐                                  │
    │  DOCX ─┤  → 文档加载 → 文本切块             │
    │  TXT ─┤       │            │              │
    │  CSV ─┘       │            ▼              │
    │               │      ┌──────────┐        │
    │               │      │ 块1: "xxx" │        │
    │               │      │ 块2: "yyy" │        │
    │               │      │ 块3: "zzz" │        │
    │               │      └─────┬────┘        │
    │               │            ▼              │
    │               │      Embedding 模型       │
    │               │      块1→[0.1,0.3,...]    │
    │               │      块2→[0.2,0.1,...]    │
    │               │            │              │
    │               │            ▼              │
    │               │      存入向量数据库        │
    └───────────────┼──────────────────────────┘
                    │
    ┌───────────────┼──────────────────────────┐
    │               │   在线阶段（每次提问）      │
    │               ▼                           │
    │         用户提问: "怎么安装?"               │
    │               │                           │
    │               ▼                           │
    │          Embedding 模型                   │
    │          → [0.15, 0.32, ...]              │
    │               │                           │
    │               ▼                           │
    │          向量库相似度检索                   │
    │          → 块2(最相关), 块5, 块1...        │
    │               │                           │
    │               ▼                           │
    │          拼接 Prompt:                      │
    │          "参考以下文档回答问题...            │
    │           [块2内容][块5内容][块1内容]        │
    │           问题: 怎么安装?"                  │
    │               │                           │
    │               ▼                           │
    │          LLM 生成回答                      │
    │          → "安装步骤如下: 1.固定墙上..."     │
    └──────────────────────────────────────────┘
""")


# ========================================
# Part 2: 统一的文档数据结构
# ========================================

print("=== Part 2: 统一数据结构 ===\n")

class LoadedDocument(BaseModel):
    """所有格式加载后都转为这个结构"""
    content: str              # 文档正文
    source: str               # 来源文件名（会传给 Chunk，最终出现在引用里）
    file_type: str            # 文件类型: pdf / docx / txt / csv
    page_count: int = 1       # 页数（PDF/Word有概念，TXT算1页）
    loaded_at: datetime = None

    def model_post_init(self, __context):
        if self.loaded_at is None:
            self.loaded_at = datetime.now()

print("所有文档加载后统一为 LoadedDocument:")
print(LoadedDocument.model_json_schema())


# ========================================
# Part 3: 文档加载器实现
# ========================================

print("\n=== Part 3: 文档加载器 ===\n")

class DocumentLoader:
    """统一的文档加载器，根据文件扩展名自动选择解析方式"""

    @staticmethod
    def load(file_path: str) -> LoadedDocument:
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return DocumentLoader._load_pdf(path)
        elif suffix == ".docx":
            return DocumentLoader._load_docx(path)
        elif suffix == ".txt":
            return DocumentLoader._load_txt(path)
        elif suffix == ".csv":
            return DocumentLoader._load_csv(path)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")

    @staticmethod
    def _load_pdf(path: Path) -> LoadedDocument:
        """用 PyMuPDF 加载 PDF"""
        doc = fitz.open(path)
        full_text = []
        for page in doc:
            full_text.append(page.get_text())
        page_count = doc.page_count       # 在 close 之前取
        doc.close()
        content = "\n".join(full_text).strip()
        return LoadedDocument(
            content=content,
            source=path.name,
            file_type="pdf",
            page_count=page_count,
        )

    @staticmethod
    def _load_docx(path: Path) -> LoadedDocument:
        """用 python-docx 加载 Word 文档"""
        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n".join(paragraphs)
        return LoadedDocument(
            content=content,
            source=path.name,
            file_type="docx",
            page_count=len(doc.paragraphs) // 30 + 1,  # 粗略估算页数
        )

    @staticmethod
    def _load_txt(path: Path) -> LoadedDocument:
        """加载纯文本"""
        content = path.read_text(encoding="utf-8")
        return LoadedDocument(
            content=content.strip(),
            source=path.name,
            file_type="txt",
        )

    @staticmethod
    def _load_csv(path: Path) -> LoadedDocument:
        """加载 CSV，拼接为可读的文本（方便后续 Embedding）"""
        with open(path, "r", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            rows = list(reader)

        if not rows:
            return LoadedDocument(content="", source=path.name, file_type="csv")

        # CSV 转成 Markdown 表格样式，方便 LLM 理解
        headers = list(rows[0].keys())
        lines = [" | ".join(headers)]
        lines.append(" | ".join(["---"] * len(headers)))
        for row in rows:
            lines.append(" | ".join(str(row[h]) for h in headers))
        content = "\n".join(lines)
        return LoadedDocument(
            content=content,
            source=path.name,
            file_type="csv"
        )


# ========================================
# Part 4: 批量加载测试
# ========================================

print("=== Part 4: 批量加载测试文档 ===\n")

test_dir = Path("E:/learn/day15/test_docs")
loader = DocumentLoader()

for file_path in sorted(test_dir.iterdir()):
    if file_path.suffix.lower() not in [".pdf", ".docx", ".txt", ".csv"]:
        continue

    print(f"加载: {file_path.name}")
    doc = loader.load(str(file_path))

    # 截取前 150 个字符展示
    preview = doc.content[:150].replace("\n", "\\n")
    if len(doc.content) > 150:
        preview += "..."

    print(f"  类型: {doc.file_type} | 页数: {doc.page_count} | 字符: {len(doc.content)}")
    print(f"  内容预览: {preview}")
    print()

print("*** Day 15 基础练习全部完成! ***")
print("\n关键结论：")
print("  1. RAG 分离线阶段（建库）和在线阶段（检索+回答）")
print("  2. 所有格式加载后统一为 LoadedDocument，后续步骤只认这一种")
print("  3. DocumentLoader.load() 根据扩展名自动路由到对应加载器")
print("  4. CSV 特殊处理为 Markdown 表格，方便 LLM 理解")
